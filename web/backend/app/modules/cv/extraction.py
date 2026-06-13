from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from pypdf import PdfReader


PARSER_VERSION = "cv-text-1.0"


class ExtractionError(ValueError):
    pass


def _normalize_text(parts: list[str]) -> str:
    normalized_parts = []
    for part in parts:
        lines = [re.sub(r"\s+", " ", line).strip() for line in part.splitlines()]
        normalized = "\n".join(line for line in lines if line)
        if normalized:
            normalized_parts.append(normalized)
    return "\n\n".join(normalized_parts).strip()


def extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        text = _normalize_text([(page.extract_text() or "") for page in reader.pages])
    except Exception as exc:
        raise ExtractionError("The PDF could not be parsed.") from exc

    if not text:
        raise ExtractionError("The PDF does not contain extractable text.")
    return text


def extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = _normalize_text(parts)
    except Exception as exc:
        raise ExtractionError("The DOCX file could not be parsed.") from exc

    if not text:
        raise ExtractionError("The DOCX file does not contain extractable text.")
    return text


def extract_text(content: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return extract_pdf_text(content)
    if file_type == "docx":
        return extract_docx_text(content)
    raise ExtractionError("No parser is available for this file type.")
